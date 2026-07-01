"""Word document export for chatbot answers.

Turns a single assistant message into a downloadable .docx file, so users
can save any answer the chatbot produces — a drafted e-mail, a memo, a
short report — as a Word document.

The chatbot's answers are written in Markdown. This module performs a
light conversion of the most common Markdown constructs into Word
formatting: headings, bullet and numbered lists, bold spans, and plain
paragraphs. It is deliberately NOT a full Markdown engine; any construct
it does not recognise is written as ordinary paragraph text, so no content
is ever lost — at worst a rare piece of formatting is flattened.

The single heavy function (build_message_docx) is cached on the message
text, so re-rendering the conversation on every Streamlit rerun does not
rebuild a document that has not changed.
"""

import re
from io import BytesIO

import streamlit as st
from docx import Document


# MIME type of a .docx file, expected by st.download_button.
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

# Line-level Markdown patterns. Each recognises one kind of block:
#   heading:  "# Title", "## Subtitle", ... up to level 6
#   bullet:   "- item", "* item", "+ item"
#   numbered: "1. item", "2) item", ...
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\d+[.)]\s+(.*)$")

# Inline **bold** spans, rendered as bold runs inside a paragraph.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Fallback file name when no title can be derived from the message.
_DEFAULT_STEM = "chatbot-answer"


def _add_inline_text(paragraph, text: str) -> None:
    """Append text to a Word paragraph, rendering **bold** as bold runs.

    Splits the text on Markdown bold markers and adds each piece as a run,
    setting run.bold on the emphasised segments. Text without any markers
    is added as a single plain run.

    Args:
        paragraph: The python-docx paragraph to append to.
        text: The inline text, possibly containing **bold** spans.
    """
    cursor = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _strip_inline_markers(text: str) -> str:
    """Remove leftover inline Markdown markers from a heading or title.

    Headings are written with python-docx's heading styles, which do not
    interpret Markdown, so any stray ** or ` markers must be dropped to
    avoid showing them literally.

    Args:
        text: The raw heading/title text.

    Returns:
        The text without bold or code markers, stripped of whitespace.
    """
    return text.replace("**", "").replace("`", "").strip()


def _render_markdown_into_document(document: Document, text: str) -> None:
    """Write a Markdown message into a Word document, block by block.

    Each non-empty line is classified as a heading, a bullet-list item, a
    numbered-list item, or a plain paragraph, and added with the matching
    Word style. Blank lines are skipped: in the chatbot's output every
    paragraph already sits on its own line, so one line maps to one Word
    block.

    Args:
        document: The python-docx Document being built.
        text: The assistant message in Markdown.
    """
    wrote_anything = False

    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(_strip_inline_markers(heading.group(2)), level=level)
            wrote_anything = True
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_text(paragraph, bullet.group(1))
            wrote_anything = True
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_text(paragraph, numbered.group(1))
            wrote_anything = True
            continue

        paragraph = document.add_paragraph()
        _add_inline_text(paragraph, line)
        wrote_anything = True

    if not wrote_anything:
        # Never save a completely empty document; add one blank paragraph
        # so the file is valid and openable.
        document.add_paragraph("")


@st.cache_data(show_spinner=False)
def build_message_docx(text: str) -> bytes:
    """Build a Word document from a chatbot answer, as raw bytes.

    Converts the Markdown answer into a .docx (headings, lists, bold and
    paragraphs) and returns the file as bytes, ready to hand to
    st.download_button. Nothing is written to disk.

    The result is cached on the exact message text, so re-rendering the
    conversation on every rerun reuses the already-built document instead
    of regenerating it.

    Args:
        text: The assistant message in Markdown. May be empty.

    Returns:
        The complete .docx file as a bytes object.
    """
    document = Document()
    _render_markdown_into_document(document, text or "")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def filename_for_message(text: str) -> str:
    """Derive a friendly .docx file name from a chatbot answer.

    Uses the first meaningful line of the message (a heading or the first
    sentence), stripped of Markdown markers and reduced to a filesystem-
    safe slug, so a downloaded resignation letter is named after its own
    first line rather than "response.docx".

    Args:
        text: The assistant message in Markdown.

    Returns:
        A safe file name ending in ".docx". Falls back to a generic name
        when no usable title can be found.
    """
    stem = _DEFAULT_STEM
    for raw_line in (text or "").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        # Drop leading block markers (#, -, *, +, "1.") then inline ones.
        line = re.sub(r"^#{1,6}\s+", "", line)
        line = re.sub(r"^[-*+]\s+", "", line)
        line = re.sub(r"^\d+[.)]\s+", "", line)
        line = _strip_inline_markers(line)
        if line:
            stem = line
            break

    # Keep only word characters, spaces and hyphens; collapse to a slug.
    safe = re.sub(r"[^\w\s-]", "", stem, flags=re.UNICODE).strip()
    safe = re.sub(r"\s+", "-", safe)[:60].strip("-")
    return f"{safe or _DEFAULT_STEM}.docx"
